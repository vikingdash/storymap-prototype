"""Unit tests for document_extractor.py — .docx validation, extraction, and citation
location lookup for the 'Analyze a company' internal-documents upload feature. No Flask,
no jobs.py, no network, no Anthropic API calls: every test here builds a real .docx
in-memory with python-docx and feeds its bytes straight to the module under test.

Run with: python3 -m unittest test_document_extractor -v
"""
import io
import unittest
import zipfile

import docx

import document_extractor as de


def _build_docx(build_fn):
    d = docx.Document()
    build_fn(d)
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


class FilenameValidation(unittest.TestCase):
    def test_docx_extension_is_accepted(self):
        de.validate_filename("report.docx")  # no exception

    def test_docx_extension_is_case_insensitive(self):
        de.validate_filename("Report.DOCX")  # no exception

    def test_wrong_extension_is_rejected(self):
        with self.assertRaises(de.UnsupportedDocumentTypeError):
            de.validate_filename("report.pdf")

    def test_macro_enabled_docm_is_rejected_even_though_never_executed(self):
        with self.assertRaises(de.UnsupportedDocumentTypeError):
            de.validate_filename("report.docm")

    def test_legacy_doc_is_rejected(self):
        with self.assertRaises(de.UnsupportedDocumentTypeError):
            de.validate_filename("report.doc")

    def test_no_extension_is_rejected(self):
        with self.assertRaises(de.UnsupportedDocumentTypeError):
            de.validate_filename("report")


class RejectionPaths(unittest.TestCase):
    """Every rejection message contains only filename/size/category — never document
    content — checked explicitly for the two paths that touch real bytes."""

    def test_empty_file_is_rejected(self):
        with self.assertRaises(de.EmptyDocumentError):
            de.validate_and_extract("empty.docx", b"")

    def test_oversized_file_is_rejected(self):
        raw = b"PK\x03\x04" + b"0" * (de.MAX_FILE_BYTES + 1)
        with self.assertRaises(de.DocumentTooLargeError):
            de.validate_and_extract("big.docx", raw)

    def test_password_protected_ole_signature_is_rejected_with_specific_message(self):
        raw = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1" + b"0" * 100
        with self.assertRaises(de.PasswordProtectedDocumentError) as ctx:
            de.validate_and_extract("secret.docx", raw)
        self.assertIn("password-protected", str(ctx.exception))
        self.assertNotIn("0000", str(ctx.exception))  # never echoes the raw bytes

    def test_random_garbage_is_rejected_as_corrupted(self):
        with self.assertRaises(de.CorruptedDocumentError):
            de.validate_and_extract("garbage.docx", b"this is not a docx file at all, just text")

    def test_zip_that_is_not_a_word_document_is_rejected(self):
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("hello.txt", "not a word document")
        with self.assertRaises(de.CorruptedDocumentError):
            de.validate_and_extract("fake.docx", buf.getvalue())

    def test_valid_docx_with_no_text_is_rejected_as_empty(self):
        raw = _build_docx(lambda d: None)
        with self.assertRaises(de.EmptyDocumentError):
            de.validate_and_extract("blank.docx", raw)

    def test_wrong_extension_is_rejected_before_any_parsing(self):
        with self.assertRaises(de.UnsupportedDocumentTypeError):
            de.validate_and_extract("notes.txt", b"hello world")

    def test_rejection_messages_never_include_document_text(self):
        raw = _build_docx(lambda d: d.add_paragraph("TOP SECRET FINANCIAL PROJECTIONS 2027"))
        # Corrupt it after the fact so it fails to parse, but keep enough of a docx shell
        # that if the extractor ever accidentally echoed content it would show up here.
        corrupted = raw[:200]
        try:
            de.validate_and_extract("secret_plan.docx", corrupted)
        except de.DocumentUploadError as exc:
            self.assertNotIn("TOP SECRET", str(exc))
            self.assertNotIn("FINANCIAL PROJECTIONS", str(exc))


class ExtractionOrderAndStructure(unittest.TestCase):
    def test_paragraphs_and_headings_preserve_document_order(self):
        def build(d):
            d.add_heading("Intro", level=1)
            d.add_paragraph("First paragraph.")
            d.add_heading("Details", level=1)
            d.add_paragraph("Second paragraph.")
        raw = _build_docx(build)
        result = de.validate_and_extract("order.docx", raw)
        lines = result["flatText"].splitlines()
        self.assertEqual(lines, ["Intro", "First paragraph.", "Details", "Second paragraph."])

    def test_table_between_two_paragraphs_stays_in_true_document_order(self):
        # The specific bug this guards against: python-docx's own document.paragraphs and
        # document.tables are two SEPARATE lists with no combined ordering — naively
        # reading paragraphs-then-tables would move this table to the end.
        def build(d):
            d.add_paragraph("Before the table.")
            t = d.add_table(rows=1, cols=2)
            t.rows[0].cells[0].text = "A"
            t.rows[0].cells[1].text = "B"
            d.add_paragraph("After the table.")
        raw = _build_docx(build)
        result = de.validate_and_extract("interleaved.docx", raw)
        lines = result["flatText"].splitlines()
        self.assertEqual(lines[0], "Before the table.")
        self.assertIn("A | B", lines[1])
        self.assertEqual(lines[2], "After the table.")

    def test_heading_levels_build_a_breadcrumb_path(self):
        def build(d):
            d.add_heading("Company Overview", level=1)
            d.add_heading("Financials", level=2)
            d.add_paragraph("Revenue was strong.")
        raw = _build_docx(build)
        result = de.validate_and_extract("headings.docx", raw)
        para_blocks = [b for b in result["structureBlocks"] if b["type"] == "paragraph"]
        self.assertEqual(para_blocks[0]["headingPath"], ["Company Overview", "Financials"])

    def test_sibling_heading_pops_the_previous_ones_own_subheadings(self):
        def build(d):
            d.add_heading("Section A", level=1)
            d.add_heading("Sub A1", level=2)
            d.add_heading("Section B", level=1)  # sibling of Section A, not nested under Sub A1
            d.add_paragraph("Under B.")
        raw = _build_docx(build)
        result = de.validate_and_extract("siblings.docx", raw)
        para_blocks = [b for b in result["structureBlocks"] if b["type"] == "paragraph"]
        self.assertEqual(para_blocks[0]["headingPath"], ["Section B"])

    def test_table_cells_are_individually_located(self):
        def build(d):
            t = d.add_table(rows=2, cols=2)
            t.rows[0].cells[0].text = "Metric"
            t.rows[0].cells[1].text = "Value"
            t.rows[1].cells[0].text = "Revenue"
            t.rows[1].cells[1].text = "$120M"
        raw = _build_docx(build)
        result = de.validate_and_extract("table.docx", raw)
        cell_blocks = [b for b in result["structureBlocks"] if b["type"] == "table_cell"]
        self.assertEqual(len(cell_blocks), 4)
        revenue_value = next(b for b in cell_blocks if b["text"] == "$120M")
        self.assertEqual(revenue_value["row"], 1)
        self.assertEqual(revenue_value["col"], 1)

    def test_extraction_is_truncated_at_the_character_cap(self):
        def build(d):
            d.add_paragraph("word " * 20000)  # far more than MAX_EXTRACTED_CHARS
        raw = _build_docx(build)
        result = de.validate_and_extract("long.docx", raw)
        self.assertTrue(result["truncated"])
        self.assertEqual(len(result["flatText"]), de.MAX_EXTRACTED_CHARS)


class LocateExcerptScope(unittest.TestCase):
    def setUp(self):
        def build(d):
            d.add_heading("Financial Performance", level=1)
            d.add_paragraph("Revenue grew twenty percent year over year.")
            t = d.add_table(rows=2, cols=2)
            t.rows[0].cells[0].text = "Metric"
            t.rows[0].cells[1].text = "Value"
            t.rows[1].cells[0].text = "Revenue"
            t.rows[1].cells[1].text = "$120M"
        raw = _build_docx(build)
        self.blocks = de.validate_and_extract("scope.docx", raw)["structureBlocks"]

    def test_locates_a_paragraph_excerpt(self):
        scope = de.locate_excerpt_scope("Revenue grew twenty percent year over year.", self.blocks)
        self.assertIn("Financial Performance", scope)
        self.assertIn("paragraph", scope)

    def test_locates_a_table_cell_excerpt(self):
        scope = de.locate_excerpt_scope("$120M", self.blocks)
        self.assertIn("Table 1", scope)
        self.assertIn("row 2", scope)

    def test_unverifiable_excerpt_degrades_to_none_rather_than_a_wrong_location(self):
        scope = de.locate_excerpt_scope("This sentence was never in the document.", self.blocks)
        self.assertIsNone(scope)

    def test_empty_excerpt_or_blocks_returns_none(self):
        self.assertIsNone(de.locate_excerpt_scope("", self.blocks))
        self.assertIsNone(de.locate_excerpt_scope("Revenue grew twenty percent year over year.", []))


class DocumentsToSources(unittest.TestCase):
    def test_shapes_match_fetch_all_sources_conventions(self):
        documents = [{
            "id": "src_upload_draft", "title": "Draft.docx", "documentRole": "current_draft_narrative",
            "flatText": "Some extracted text.", "structureBlocks": [], "retrievedAt": "2026-01-01T00:00:00Z",
        }]
        sources, source_text_by_id = de.documents_to_sources(documents)
        self.assertEqual(len(sources), 1)
        self.assertEqual(sources[0]["sourceType"], "internal")
        self.assertEqual(sources[0]["documentRole"], "current_draft_narrative")
        self.assertEqual(sources[0]["id"], "src_upload_draft")
        self.assertNotIn("flatText", sources[0])  # raw text lives only in source_text_by_id, matching fetch_all_sources
        self.assertEqual(source_text_by_id["src_upload_draft"], "Some extracted text.")


class SlugifyFilename(unittest.TestCase):
    def test_produces_a_safe_lowercase_slug(self):
        self.assertEqual(de.slugify_filename("HPS Draft Narrative (v2).docx"), "hps_draft_narrative_v2")

    def test_empty_or_extensionless_name_falls_back(self):
        self.assertEqual(de.slugify_filename(""), "document")


if __name__ == "__main__":
    unittest.main()
