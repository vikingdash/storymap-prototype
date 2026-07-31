"""Local .docx extraction for the 'Analyze a company' internal-documents upload feature.

Mirrors extractor.py's role (readable-text extraction) but for uploaded Word documents
instead of fetched HTML. Uses python-docx, which only ever reads text out of a .docx's
XML parts — it cannot execute a macro and cannot open the OLE/CFBF container Word uses
for password-protected files, which is exactly what makes both "no code execution risk"
and "clean password-protected detection" true for free rather than something this module
has to build.

Two outputs per document:
  - flatText: paragraphs and table rows joined in true document order, fed to the model
    exactly like a fetched web page's extracted text (pipeline_runner.py never knows the
    difference) and checked against by citation_verify.excerpt_is_verified() unchanged.
  - structureBlocks: an ordered list of {type, text, ...location fields} — backend-only,
    never sent to the model or the browser, used solely by locate_excerpt_scope() to turn
    an already-verified excerpt into a human-readable "Section X — paragraph N" citation
    (jobs.py's _dataset_from_checkpoint, read-time, mirroring how usage/stageProgress/
    sourceCoverage are already derived fresh on every read rather than cached).

Every rejection raises a DocumentUploadError subclass whose message contains only the
filename, byte size, and failure category — never a fragment of the document's actual
text. Every log/print statement in this module follows the same rule, matching the "never
log document contents" requirement.
"""
import io
import re
import zipfile

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from citation_verify import excerpt_is_verified

MAX_FILE_BYTES = 20 * 1024 * 1024  # 20MB per file
MAX_EXTRACTED_CHARS = 50000  # per document — bounds both prompt size and cost

ALLOWED_EXTENSION = ".docx"
# Rejected even though this app would never execute a macro either way — a mislabeled
# risky file should never even reach the parser, and a plain .doc (legacy binary format)
# isn't OOXML at all, so python-docx can't read it regardless.
DISALLOWED_EXTENSIONS = {".docm", ".dotm", ".dot", ".doc"}

_OLE_MAGIC = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"  # password-protected/legacy Office signature
_ZIP_MAGIC = b"PK\x03\x04"

_HEADING_STYLE_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


class DocumentUploadError(Exception):
    """Base for every rejection — app.py catches this one class and returns str(exc) as
    the client-facing error message. Never constructed with document content, only
    filename/size/category — see this module's docstring."""


class UnsupportedDocumentTypeError(DocumentUploadError):
    pass


class DocumentTooLargeError(DocumentUploadError):
    pass


class PasswordProtectedDocumentError(DocumentUploadError):
    pass


class CorruptedDocumentError(DocumentUploadError):
    pass


class EmptyDocumentError(DocumentUploadError):
    pass


def _extension_of(filename):
    lower = (filename or "").lower()
    return lower[lower.rfind("."):] if "." in lower else ""


def validate_filename(filename):
    """Extension check only — a filename can lie about content, which is why
    validate_and_extract() ALSO sniffs the real bytes below; this check exists so an
    obviously-wrong file is rejected with a specific, fast message before any parsing."""
    ext = _extension_of(filename)
    if ext in DISALLOWED_EXTENSIONS:
        raise UnsupportedDocumentTypeError(
            f'"{filename}": macro-enabled or legacy Word formats are not supported — only plain .docx.'
        )
    if ext != ALLOWED_EXTENSION:
        raise UnsupportedDocumentTypeError(
            f'"{filename}": only .docx files are supported (got "{ext or "no extension"}").'
        )


def slugify_filename(filename):
    stem = filename.rsplit(".", 1)[0] if filename and "." in filename else (filename or "document")
    slug = re.sub(r"[^a-z0-9]+", "_", stem.lower()).strip("_")
    return (slug or "document")[:50]


def _iter_block_items(document):
    """Yields Paragraph/Table objects in TRUE document order. python-docx's own
    document.paragraphs and document.tables are two separate lists with no combined
    ordering between them — reading one after the other would silently misorder any
    document where a table sits between two paragraphs. Walking the underlying XML body
    directly (dispatching on each child's own tag) is the standard way to avoid that."""
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def _walk_body(document):
    """Returns (flat_lines, structure_blocks). Headings are detected from Word's built-in
    style names ("Heading 1".."Heading 9", "Title") — a heuristic that only recognizes
    documents using Word's standard heading styles, not manually-bolded text standing in
    for a heading. Tables are read cell by cell; merged cells and nested tables are
    best-effort (python-docx exposes a merged cell's text at every position it spans
    rather than modeling the visual merge), matching "tables where practical.\""""
    flat_lines = []
    blocks = []
    heading_path = []  # stack of (level, text)
    para_index = 0
    table_index = 0

    for item in _iter_block_items(document):
        if isinstance(item, Paragraph):
            text = item.text.strip()
            if not text:
                continue
            style_name = (item.style.name if item.style else "") or ""
            heading_match = _HEADING_STYLE_RE.match(style_name)
            is_title = style_name.strip().lower() == "title"
            if heading_match or is_title:
                level = 0 if is_title else int(heading_match.group(1))
                while heading_path and heading_path[-1][0] >= level:
                    heading_path.pop()
                heading_path.append((level, text))
                flat_lines.append(text)
                blocks.append({"type": "heading", "level": level, "text": text})
            else:
                para_index += 1
                flat_lines.append(text)
                blocks.append({
                    "type": "paragraph", "index": para_index,
                    "headingPath": [t for _, t in heading_path], "text": text,
                })
        else:
            table_index += 1
            for row_idx, row in enumerate(item.rows):
                cell_texts = [c.text.strip() for c in row.cells]
                if not any(cell_texts):
                    continue
                flat_lines.append(f"[Table {table_index}] " + " | ".join(cell_texts))
                for col_idx, cell_text in enumerate(cell_texts):
                    if not cell_text:
                        continue
                    blocks.append({
                        "type": "table_cell", "tableIndex": table_index, "row": row_idx, "col": col_idx,
                        "headingPath": [t for _, t in heading_path], "text": cell_text,
                    })
    return flat_lines, blocks


def validate_and_extract(filename, raw_bytes):
    """Runs every rejection check in order, then extracts. Returns
    {flatText, structureBlocks, wordCount, truncated}. Raises a DocumentUploadError
    subclass on any failure, filename/size/category only — never document content."""
    validate_filename(filename)

    if not raw_bytes:
        raise EmptyDocumentError(f'"{filename}" is empty (0 bytes).')
    if len(raw_bytes) > MAX_FILE_BYTES:
        raise DocumentTooLargeError(
            f'"{filename}" is {len(raw_bytes)} bytes, exceeding the {MAX_FILE_BYTES}-byte ({MAX_FILE_BYTES // (1024 * 1024)}MB) limit.'
        )

    # Password-protected .docx isn't a zip at all (Word wraps it in an OLE/CFBF
    # container instead) — checked before ever handing the bytes to python-docx, so this
    # gets a specific, accurate message instead of a generic "corrupted" one.
    if raw_bytes[:8] == _OLE_MAGIC:
        raise PasswordProtectedDocumentError(
            f'"{filename}" appears to be password-protected (or an old .doc file) — it cannot be read.'
        )
    if raw_bytes[:4] != _ZIP_MAGIC:
        raise CorruptedDocumentError(f'"{filename}" is not a valid .docx file.')

    try:
        with zipfile.ZipFile(io.BytesIO(raw_bytes)) as zf:
            if "word/document.xml" not in zf.namelist():
                raise CorruptedDocumentError(
                    f'"{filename}" is a zip file but not a Word document (missing word/document.xml).'
                )
    except zipfile.BadZipFile as exc:
        raise CorruptedDocumentError(f'"{filename}" could not be opened as a .docx file.') from exc

    try:
        document = docx.Document(io.BytesIO(raw_bytes))
        flat_lines, blocks = _walk_body(document)
    except DocumentUploadError:
        raise
    except Exception as exc:  # python-docx/lxml can raise several distinct exception types for malformed XML
        raise CorruptedDocumentError(f'"{filename}" could not be parsed as a .docx file.') from exc

    flat_text = "\n".join(flat_lines).strip()
    word_count = len(flat_text.split())
    if word_count == 0:
        raise EmptyDocumentError(f'"{filename}" has no extractable text.')

    truncated = len(flat_text) > MAX_EXTRACTED_CHARS
    if truncated:
        flat_text = flat_text[:MAX_EXTRACTED_CHARS]

    return {"flatText": flat_text, "structureBlocks": blocks, "wordCount": word_count, "truncated": truncated}


def documents_to_sources(documents):
    """documents: the persisted checkpoint["uploadedDocuments"]["documents"] list — each
    entry already carries id/title/documentRole/flatText/structureBlocks/retrievedAt from
    validate_and_extract() + role-tagging at upload time (app.py). Returns (sources,
    source_text_by_id) shaped EXACTLY like pipeline_runner.fetch_all_sources()'s own
    return values, so jobs.py/pipeline_runner.py can merge the two source lists and treat
    every source uniformly from that point on — neither the prompt-building code nor the
    evidence/citation/confidence machinery needs to know a source came from a file rather
    than a URL."""
    sources = []
    source_text_by_id = {}
    for doc in documents:
        sources.append({
            "id": doc["id"],
            "companyId": "live",
            "title": doc["title"],
            "publisher": "Internal upload",
            "sourceType": "internal",
            "documentRole": doc["documentRole"],
            "retrievedAt": doc["retrievedAt"],
            "permissionStatus": "approved",
        })
        source_text_by_id[doc["id"]] = doc["flatText"]
    return sources, source_text_by_id


def locate_excerpt_scope(excerpt, structure_blocks):
    """Given a model-produced excerpt that has ALREADY been verified as a real substring
    of the document's flat text (citation_verify.excerpt_is_verified — never re-decided
    here) and that document's structure map, returns a short human-readable location
    ("Section 'X' — paragraph 4", "Table 2, row 3"), or None if no single block's own
    text contains the excerpt (e.g. an excerpt spanning two blocks after paragraph-
    boundary joining — this degrades to no location rather than guessing a wrong one)."""
    if not excerpt or not structure_blocks:
        return None
    for block in structure_blocks:
        if block["type"] == "heading":
            continue
        if excerpt_is_verified(excerpt, block["text"]):
            return _describe_location(block)
    return None


def _describe_location(block):
    path = " > ".join(block.get("headingPath") or []) or None
    if block["type"] == "paragraph":
        loc = f"paragraph {block['index']}"
    else:
        loc = f"Table {block['tableIndex']}, row {block['row'] + 1}"
    return f"{path} — {loc}" if path else loc[0].upper() + loc[1:]
